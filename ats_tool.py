from fastapi import Request, UploadFile, File, Form
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
import re
from typing import List, Optional
from database import get_user_by_id, set_user_credits

templates = Jinja2Templates(directory="templates")

# High-risk measures that REQUIRE mandatory airtightness testing
HIGH_RISK_MEASURES = {'internal_wall', 'external_wall', 'cavity_wall'}

# Complete measure descriptions
MEASURE_DESCRIPTIONS = {
    'loft_insulation': {
        'name': 'Loft insulation top-up',
        'impact': 'Installation of additional loft insulation will require work in the roof space, creating potential leakage paths through the ceiling if not properly sealed.',
        'controls': '''Loft Access and Penetrations:
- Seal all gaps around loft hatch frame using expanding foam or flexible sealant before insulation installation
- Install compression seals or draught-proofing strips to loft hatch perimeter
- Ensure loft hatch is fitted with catches to maintain seal when closed
- Seal all service penetrations through ceiling (electrical cables, pipework) with appropriate fire-rated materials
- Install proprietary cable grommets or foam collars around penetrations
- Ensure insulation does not compress existing air sealing measures at eaves or party walls
- Maintain minimum 50mm clearance to recessed light fittings unless fire-rated hoods installed
- Photograph all sealed areas before covering with insulation''',
        'verification': '''Loft Insulation Airtightness Verification:
- Visual inspection of all ceiling penetrations and loft hatch sealing before insulation installation
- Post-installation smoke pencil test around loft hatch perimeter to verify seal
- Thermal imaging of ceiling to identify any unsealed areas (if available)
- Verification that all service penetrations remain sealed after insulation works
- Check that loft hatch compression seal makes full contact around entire perimeter'''
    },
    'cavity_wall': {
        'name': 'Cavity wall insulation (CWI)',
        'impact': 'Cavity wall insulation installation requires multiple penetrations through the external wall which must be properly sealed to prevent air leakage paths. This is a HIGH-RISK measure requiring mandatory post-installation airtightness testing.',
        'controls': '''Cavity Wall Installation Airtightness Measures:
- Mark all injection hole locations and photograph before drilling
- Use appropriate drill bit size to minimize hole diameter (typically 22mm)
- Ensure cavity is properly filled without voids, particularly around window reveals and wall junctions
- Seal all injection holes with matching mortar mix, ensuring full depth fill and flush finish
- For older properties with uncertain cavity construction, conduct trial hole inspection first
- Verify cavity ties and wall construction before proceeding
- Inspect for and seal any existing cracks in mortar before installation
- Particular attention to sealing around window and door reveals where cavity closes''',
        'verification': '''Cavity Wall Insulation Verification:
- Borescope inspection of sample cavities to verify complete fill and no voids
- Visual inspection of all filled injection holes to ensure proper sealing
- Thermographic survey post-installation to identify any unfilled areas or thermal bridging
- Verification that all mortar repairs are fully cured and sealed
- Check for and seal any cracks that may have developed during installation process
- Smoke pencil test around window and door reveals to verify seal integrity'''
    },
    'external_wall': {
        'name': 'External wall insulation (EWI)',
        'impact': 'EWI system creates a new external envelope. Critical junctions at windows, doors, roof, and ground level must maintain airtightness. This is a HIGH-RISK measure requiring mandatory post-installation airtightness testing.',
        'controls': '''External Wall Insulation Airtightness Strategy:
- Ensure continuity of airtightness layer at all junctions before EWI installation
- Seal all existing cracks and gaps in substrate wall using appropriate filler
- Install perimeter sealing tape at window and door reveals before EWI boards
- Use flexible sealant at all window/door frame to EWI system junctions
- Ensure proper lapping and sealing of EWI base track to prevent air infiltration
- Install cavity closers or other appropriate details at reveals to maintain air barrier continuity
- Seal ventilator penetrations and ensure they maintain weathertightness
- Service penetrations (meter boxes, vents) to be sealed with appropriate gaskets/seals
- Render finish to be continuous and crack-free to maintain weather resistance''',
        'verification': '''External Wall Insulation Verification:
- Pre-installation survey and sealing of all substrate cracks and gaps
- Visual inspection of all junctions and penetrations during installation
- Thermographic survey to identify any gaps in insulation or air leakage paths
- Smoke pencil testing at all internal reveals and junctions
- Post-installation air permeability test to verify improvement
- Check all service penetrations for proper sealing and weatherproofing
- Verify render finish is complete, crack-free, and properly cured'''
    },
    'internal_wall': {
        'name': 'Internal wall insulation (IWI)',
        'impact': 'IWI creates junctions with existing building fabric. Proper detailing required at floor, ceiling, and party wall junctions to maintain airtightness. This is a HIGH-RISK measure requiring mandatory post-installation airtightness testing.',
        'controls': '''Internal Wall Insulation Airtightness Measures:
- Seal all gaps, cracks and penetrations in existing wall before IWI installation
- Install continuous air barrier (typically vapor control layer) as part of IWI system
- Tape or seal all joints and overlaps in vapor control layer
- Seal perimeter of VCL to floor, ceiling, and returns using appropriate tape or sealant
- Particular attention to maintaining air barrier continuity at party wall junctions
- Seal all service penetrations (electrical boxes, pipes) through IWI with gaskets or sealant
- Install proprietary airtight electrical boxes where possible
- Ensure plasterboard joints are properly filled and sealed
- Seal skirting board to floor junction with decorators caulk''',
        'verification': '''Internal Wall Insulation Verification:
- Visual inspection of vapor control layer installation and taping before boarding
- Smoke pencil test at all electrical penetrations and service entries
- Inspection of all junctions (floor/wall/ceiling) for continuity of air barrier
- Post-installation air permeability test
- Thermal imaging to identify any gaps in insulation or thermal bridging
- Verification that all plasterboard joints are filled and smooth
- Check for any defects in VCL that may have occurred during installation'''
    },
    'floor_insulation': {
        'name': 'Floor insulation',
        'impact': 'Floor insulation installation may disturb existing skirting boards and floor junctions. Suspended timber floors require particular attention to maintain air barrier.',
        'controls': '''Floor Insulation Airtightness Strategy:
- For suspended timber floors: ensure underfloor void remains ventilated to prevent condensation
- Seal gaps between floorboards using appropriate filler before overlay (if applicable)
- Install breather membrane over insulation ensuring continuity and proper overlap/taping
- For solid floors: seal any cracks in existing floor slab before insulation installation
- Seal perimeter junction between floor insulation and wall using flexible sealant
- Reinstall skirting boards with bead of sealant at floor junction
- Seal all service penetrations through floor (pipes, cables) with appropriate materials
- Ensure underfloor ventilation (airbricks) remain clear and functional
- For party walls: maintain fire stopping and air sealing at floor/wall junction''',
        'verification': '''Floor Insulation Verification:
- Visual inspection of all floor perimeter sealing before skirting installation
- Verification that underfloor ventilation airbricks are clear and functional (suspended floors)
- Check all service penetrations for proper sealing
- Smoke pencil test at skirting board/floor junction
- Thermal imaging of floor to identify any gaps or cold spots
- Post-installation air permeability test
- Verify breather membrane (if used) is continuous with no tears or gaps'''
    },
    'solar_pv': {
        'name': 'Solar PV installation',
        'impact': 'Solar PV installation requires roof penetrations for fixings and cable entry which must be properly weatherproofed and sealed to maintain airtightness.',
        'controls': '''Solar PV Installation Airtightness Measures:
- All roof penetrations for PV fixings to be sealed with appropriate weatherproof sealant
- Use proprietary mounting systems with integral sealing washers
- Cable entry points through roof membrane to be sealed with grommets and waterproof sealant
- Consider using external conduit routing where possible to minimize roof penetrations
- Ensure roofing underlay integrity is maintained around all penetrations
- If passing through loft space, seal cable entry with fire-rated materials
- Maintain minimum clearance to roof covering to allow ventilation
- Verify all flashings and weatherproofing details are correctly installed
- DC isolator and cabling through building envelope to be sealed with appropriate materials''',
        'verification': '''Solar PV Airtightness Verification:
- Visual inspection of all roof penetrations for proper sealing and weatherproofing
- Internal inspection of cable entry points through roof/wall for air sealing
- Verification that roofing underlay has not been damaged or compromised
- Check all external junction boxes for weatherproof gaskets and seals
- Smoke pencil test at internal cable entry points (where accessible)
- Thermal imaging to verify no air leakage around penetrations
- Post-installation inspection during rainfall to verify weathertightness'''
    },
    'heating_upgrade': {
        'name': 'Heating system upgrade',
        'impact': 'New heating system may require flue penetrations, pipework, and electrical connections through building fabric.',
        'controls': '''Heating System Upgrade Airtightness Strategy:
- Seal all new flue penetrations through walls/roof with proprietary flashing and fire-rated sealing
- For balanced flue systems, ensure wall plate is sealed to wall with appropriate gasket/sealant
- Seal all pipework penetrations through walls, floors, and ceilings
- Use proprietary pipe sleeves or collars with flexible sealant
- For underfloor heating: seal all floor penetrations for pipework and manifold connections
- Ensure any old flue penetrations from removed system are properly sealed and made good
- Radiator pipe penetrations through floors to be sealed with acoustic sealant
- Gas supply pipe penetrations to be sealed appropriately
- Electrical supply penetrations for boiler/controls to use grommets or gaskets''',
        'verification': '''Heating System Verification:
- Visual inspection of all new flue penetrations for proper sealing
- Verification of balanced flue wall plate seal
- Smoke pencil test at all accessible pipe penetrations
- Check that redundant flue penetrations from old system are properly sealed
- Thermal imaging around new flue penetration to verify proper insulation
- Inspection of all radiator pipe floor penetrations
- Post-installation air permeability test
- Gas Safe certification that all gas penetrations meet regulations'''
    },
    'windows_doors': {
        'name': 'Windows and doors replacement',
        'impact': 'Window and door replacement is a major intervention affecting airtightness. Proper installation and sealing is critical.',
        'controls': '''Windows and Doors Installation Airtightness Strategy:
- Remove old units carefully to avoid damage to reveals
- Clean and prepare all reveals, seal any cracks or defects in substrate
- Install new units in accordance with manufacturer instructions
- Use appropriate fixings with sealing washers
- Install perimeter seal between frame and reveal (polyurethane foam or expanding tape)
- Apply internal sealant bead between frame and reveal after foam has cured
- Ensure cavity closers or insulated reveals maintain thermal and air barrier continuity
- Install sill flashing and weatherbar correctly
- External mastic seal to frame perimeter for weather resistance
- Verify trickle ventilators (if fitted) have functional controls and seals
- Install compression seals to opening sections correctly
- Check and adjust all opening lights for proper closure and seal compression''',
        'verification': '''Windows and Doors Installation Verification:
- Visual inspection of all frame to reveal sealing (internal and external)
- Smoke pencil test at frame perimeter (internal side) on sample of installations
- Verification that all opening lights close properly and compress seals
- Check operation of all handles, locks, and trickle ventilators
- Thermal imaging survey of all installations to identify gaps or thermal bridging
- Post-installation air permeability test (significant improvement expected)
- Water spray test on sample installations to verify weathertightness
- Verify U-values and specifications match specification
- Check building control completion certificate issued'''
    },
    'ventilation': {
        'name': 'Ventilation system upgrade',
        'impact': 'New ventilation system requires ductwork penetrations through building envelope and may include MVHR unit installation.',
        'controls': '''Ventilation System Installation Airtightness Strategy:
- All ductwork penetrations through external walls/roof to be sealed with proprietary sealing grommets
- Use flexible sealant around duct sleeve perimeter (both sides of wall)
- For MVHR: ensure unit is mounted on vibration-isolating pads to prevent noise transmission
- All ductwork joints to be sealed with aluminum tape or mastic
- Maintain continuity of thermal insulation around ductwork penetrations
- Install weatherproof external grilles with integral back-draught shutters
- Seal internal register mounting boxes to ceiling/wall
- Ensure any redundant extract fan penetrations are properly sealed
- For continuous mechanical extract: seal redundant trickle vent openings if required
- Commission system and verify flow rates meet design specification''',
        'verification': '''Ventilation System Verification:
- Visual inspection of all ductwork penetrations for proper sealing
- Smoke pencil test at all register locations and duct penetrations
- Commissioning test to verify designed airflow rates achieved
- Verification of external grilles for weathertightness and back-draught shutter function
- Check that redundant ventilation openings are properly sealed
- Post-installation air permeability test
- Thermal imaging of duct penetrations to verify no thermal bridging
- Verification that MVHR heat recovery efficiency meets specification (if applicable)
- Sound level testing to verify acoustic performance'''
    },
    'draught_proofing': {
        'name': 'Draught proofing',
        'impact': 'Draught proofing directly improves airtightness by sealing intentional and unintentional gaps.',
        'controls': '''Draught Proofing Installation Strategy:
- Install compression seals to all external doors (head, jambs, and threshold)
- Fit brush or blade seals to bottom of doors
- Install perimeter seals to loft hatches with compression fit
- Apply self-adhesive foam strips to window opening lights (where not already fitted)
- Seal letterbox openings with brush seals or flaps
- Seal gaps around service entry points (meter boxes, waste pipes, overflow pipes)
- Install chimney balloon or cap to unused chimneys (ensure permanent ventilation maintained if required)
- Seal gaps in suspended timber floors using appropriate sealant or strips
- Draught-proof internal doors where appropriate to control airflow
- Seal any gaps at skirting board/floor junction with decorators caulk''',
        'verification': '''Draught Proofing Verification:
- Functional check of all door and window seals for proper compression
- Visual inspection of all sealed areas
- Smoke pencil test at common leakage areas (letterbox, loft hatch, floor/wall junctions)
- Pre and post air permeability test to quantify improvement
- Verification that all seals are correctly installed and functional
- Check that ventilation strategy remains appropriate after improvements
- Thermal imaging to identify any remaining significant air leakage paths
- Verify unused chimneys are safely and appropriately sealed'''
    }
}

def generate_measures_text(selected_measures):
    measure_names = [MEASURE_DESCRIPTIONS[m]['name'] for m in selected_measures if m in MEASURE_DESCRIPTIONS]
    if not measure_names:
        return "Loft insulation; Solar PV"
    return "; ".join(measure_names)

def generate_impact_text(selected_measures):
    if not selected_measures:
        return "The proposed measures involve penetrations through the building envelope which may affect airtightness."
    impacts = [MEASURE_DESCRIPTIONS[m]['impact'] for m in selected_measures if m in MEASURE_DESCRIPTIONS]
    return " ".join(impacts)

def generate_control_measures(selected_measures):
    if not selected_measures:
        return "General airtightness measures to be implemented during installation."
    controls = []
    for i, measure in enumerate(selected_measures, 1):
        if measure in MEASURE_DESCRIPTIONS:
            controls.append(f"{i}. {MEASURE_DESCRIPTIONS[measure]['controls']}")
    controls.append(f"\n{len(controls)+1}. General Workmanship:\n• All installers briefed on airtightness requirements\n• All penetrations documented and photographed\n• Materials specification maintained throughout\n• Site inspections at key stages")
    return "\n\n".join(controls)

def generate_verification_text(selected_measures, has_high_risk):
    if not selected_measures:
        return "Post-installation air permeability testing recommended per PAS 2035."
    
    verifications = []
    for measure in selected_measures:
        if measure in MEASURE_DESCRIPTIONS:
            verifications.append(MEASURE_DESCRIPTIONS[measure]['verification'])
    
    # Add mandatory or recommended testing based on risk
    if has_high_risk:
        overall_test = '''\n⚠️ MANDATORY AIRTIGHTNESS TESTING REQUIRED ⚠️

Due to the inclusion of high-risk measures (IWI, EWI, RIR, or CWI), post-installation airtightness testing is MANDATORY per PAS 2035:2023 Annex 8.2.35.

Testing Requirements:
- Testing Standard: ATTMA Technical Standard L1 (Measuring Air Permeability in the Existing Housing Stock)
- Target Air Permeability: ≤ 10 m³/(h.m²) at 50Pa
- Testing to be conducted by ATTMA-certified air tightness tester
- Test to be performed following completion of all retrofit measures
- Results MUST be documented and sent to Brian McKevitt (Retrofit Coordinator)
- Test results to be included in project completion pack and submitted to funding body

Failure to conduct and submit airtightness test results will result in non-compliance with PAS 2035 requirements.'''
    else:
        overall_test = '''\nAirtightness Testing - RECOMMENDED:

While not mandatory for the selected measures, post-installation airtightness testing is recommended to:
- Verify effectiveness of air sealing measures
- Quantify improvement in building performance
- Provide baseline for future work

If testing is conducted:
- Testing Standard: ATTMA Technical Standard L1
- Target: ≤ 10 m³/(h.m²) at 50Pa
- Testing by ATTMA-certified tester
- Results to be documented in project completion pack'''
    
    verifications.append(overall_test)
    return "\n\n".join(verifications)

def create_ats_document(data):
    doc = Document()
    
    # Set default font to Calibri 11pt
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title - Bold, Calibri 18pt, Centered
    title = doc.add_paragraph()
    title_run = title.add_run('AIRTIGHTNESS STRATEGY')
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle - Calibri 12pt, Centered
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('PAS 2035:2023 – Annex 8.2.35 Retrofit Assessment')
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Section 1 - Property Details
    heading = doc.add_heading('1. Property Details', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Address: ').bold = True
    p.add_run(data.get('address', '[Property Address]'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Property Type: ').bold = True
    p.add_run(data.get('property_type', 'Detached house'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Construction Type: ').bold = True
    p.add_run(data.get('construction', 'Cavity wall construction'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Age/Period: ').bold = True
    p.add_run(data.get('age', 'Post-1990s'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 2 - Assessment Details
    heading = doc.add_heading('2. Assessment Details', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Assessor: ').bold = True
    p.add_run(data.get('assessor', '[Assessor Name]'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Inspection Date: ').bold = True
    p.add_run(data.get('inspection_date', '[Date]'))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Retrofit Coordinator: ').bold = True
    p.add_run('Brian McKevitt MCIOB')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('PAS 2035 Compliance: ').bold = True
    p.add_run('This Airtightness Strategy is prepared in accordance with PAS 2035:2023 Annex 8.2.35')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 3 - Proposed Measures
    heading = doc.add_heading('3. Proposed Retrofit Measures', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(data.get('measures_text', ''))
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Impact on Airtightness: ').bold = True
    p.add_run(data.get('impact_text', ''))
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add high-risk warning if applicable
    if data.get('has_high_risk'):
        p = doc.add_paragraph()
        p.add_run('\n⚠️ HIGH-RISK MEASURES SELECTED: ').bold = True
        p.add_run('This project includes high-risk measures (IWI, EWI, or CWI) which require MANDATORY post-installation airtightness testing per PAS 2035:2023.')
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(192, 0, 0)  # Red for warning
    
    doc.add_paragraph()
    
    # Section 4 - Existing Condition
    heading = doc.add_heading('4. Existing Airtightness Condition', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Baseline Assessment:').bold = True
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    existing_text = data.get('existing_condition', '''• No prior airtightness test available; typical air permeability estimated at 10-15 m³/(h.m²) at 50Pa
- Likely leakage paths identified: loft hatch perimeter, eaves junctions, service penetrations
- No persistent damp or mould observed during visual inspection
- Existing background ventilation and mechanical extract ventilation present''')
    
    p = doc.add_paragraph()
    run = p.add_run(existing_text)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 5 - Control Measures
    heading = doc.add_heading('5. Control Measures During Retrofit', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run('Air Sealing Protocols (PAS 2035:2023 Clause 8.2.35.2):').bold = True
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(data.get('control_measures', ''))
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 6 - Verification
    heading = doc.add_heading('6. Verification & Testing', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(data.get('verification', ''))
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 7 - Residual Risks
    heading = doc.add_heading('7. Residual Risks & Management', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    risks_text = '''Ventilation Adequacy (PAS 2035:2023 Clause 7.3):
- Risk: Improved airtightness without adequate ventilation may lead to condensation and indoor air quality issues
- Mitigation: Existing mechanical extract ventilation to remain operational; trickle ventilators maintained in working order
- Occupant advised to use extract fans when cooking/bathing and maintain trickle ventilator operation
- Consider future upgrade to MVHR if further airtightness improvements planned

Condensation Risk:
- Risk: Reduced air change rate may increase humidity levels
- Mitigation: Occupant guidance provided on ventilation practices; existing ventilation systems maintained
- Monitor for condensation in first heating season; respond to any issues promptly

Review and Monitoring:
- Initial review 3-6 months post-installation
- Annual review of building performance for 2 years
- Occupant feedback mechanism established
- Any defects to be addressed promptly in accordance with warranty provisions'''
    
    p = doc.add_paragraph()
    run = p.add_run(risks_text)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Section 8 - Regulatory References
    heading = doc.add_heading('8. Regulatory References & Standards', level=2)
    heading.runs[0].font.name = 'Calibri'
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    references_text = '''This Airtightness Strategy complies with:

- PAS 2035:2023 – Retrofitting dwellings for improved energy efficiency – Specification and guidance (Annex 8.2.35)
- PAS 2030:2019 – Improving the energy efficiency of existing buildings – Specification and guidance
- Building Regulations Approved Document L1B (2021 edition) – Conservation of fuel and power in existing dwellings
- Building Regulations Approved Document F (2021 edition) – Ventilation
- ATTMA Technical Standard L1 – Measuring Air Permeability in the Existing Housing Stock
- BS EN 13829:2001 – Thermal performance of buildings – Determination of air permeability of buildings
- CIBSE TM23 (2000) – Testing buildings for air leakage'''
    
    p = doc.add_paragraph()
    run = p.add_run(references_text)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer with signature
    p = doc.add_paragraph()
    p.add_run('Document Prepared By:').bold = True
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Retrofit Coordinator: Brian McKevitt MCIOB")
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Document reference
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Document Reference: ATS-{datetime.now().strftime('%Y%m%d')}\nGenerated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

async def parse_pdf(pdf_file):
    parsed = {}
    try:
        content = await pdf_file.read()
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = "\n".join([p.extract_text() or "" for p in pdf.pages])
            
            print(f"📄 PARSING: {pdf_file.filename}")
            
            # Extract assessor
            assessor_patterns = [
                r'Assessor\s+name\s+(?:Mr\.|Mrs\.|Ms\.|Dr\.)?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
                r'Name:\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\s+Title:'
            ]
            for pattern in assessor_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    parsed['assessor'] = match.group(1).strip()
                    print(f"   ✅ Assessor: {parsed['assessor']}")
                    break
            
            # Extract date
            date_match = re.search(r'Inspection\s+[Dd]ate\s+([0-9]{1,2}/[0-9]{1,2}/[0-9]{4})', text)
            if date_match:
                parsed['inspection_date'] = date_match.group(1)
                print(f"   ✅ Date: {parsed['inspection_date']}")
            
            # Extract property type
            if 'Bungalow' in text:
                if 'Semi-Detached' in text or 'Semi-detached' in text:
                    parsed['property_type'] = 'Semi-detached bungalow'
                else:
                    parsed['property_type'] = 'Bungalow'
                print(f"   ✅ Property: {parsed['property_type']}")
            elif 'Semi-Detached' in text or 'Semi-detached' in text:
                parsed['property_type'] = 'Semi-detached house'
                print(f"   ✅ Property: {parsed['property_type']}")
            
            # Extract construction
            if 'Timber Frame' in text or 'Timber frame' in text:
                parsed['construction'] = 'Timber framed construction'
                print(f"   ✅ Construction: {parsed['construction']}")
            elif 'Cavity' in text and 'wall' in text.lower():
                parsed['construction'] = 'Cavity wall construction'
                print(f"   ✅ Construction: {parsed['construction']}")
                
    except Exception as e:
        print(f"   ❌ Error: {e}")
    
    return parsed


async def ats_generator_route(request: Request, user_row):
    if request.method == "GET":
        user_id = user_row.get("id")
        user_data = get_user_by_id(user_id)
        balance = float(user_data.get("credits", 0.0))
        return templates.TemplateResponse("ats_generator.html", {"request": request, "balance": balance})
    
    elif request.method == "POST":
        user_id = user_row.get("id")
        user_data = get_user_by_id(user_id)
        balance = float(user_data.get("credits", 0.0))
        
        if balance < 10.00:
            return HTMLResponse(content="<h1>Insufficient balance</h1>", status_code=400)
        
        form_data = await request.form()
        address = form_data.get('address')
        if not address:
            return HTMLResponse(content="<h1>Please provide address</h1>", status_code=400)
        
        selected_measures = form_data.getlist('measures')
        if not selected_measures:
            return HTMLResponse(content="<h1>Please select at least one measure</h1>", status_code=400)
        
        # Check if high-risk measures selected
        has_high_risk = bool(set(selected_measures) & HIGH_RISK_MEASURES)
        
        cr_file = form_data.get("cr_file")
        if not cr_file or not hasattr(cr_file, 'read'):
            return HTMLResponse(content="<h1>Please upload a PDF</h1>", status_code=400)
        
        print("📄 Parsing Condition Report...")
        cr_data = await parse_pdf(cr_file)
        
        site_notes_file = form_data.get("site_notes_file")
        site_notes_data = {}
        if site_notes_file and hasattr(site_notes_file, 'read'):
            print("📋 Parsing Site Notes...")
            site_notes_data = await parse_pdf(site_notes_file)
        
        merged_data = {**cr_data, **site_notes_data}
        
        data = {
            'address': address,
            'property_type': merged_data.get('property_type', 'Detached house'),
            'construction': merged_data.get('construction', 'Cavity wall construction'),
            'assessor': merged_data.get('assessor', '[Assessor]'),
            'inspection_date': merged_data.get('inspection_date', '[Date]'),
            'measures_text': generate_measures_text(selected_measures),
            'impact_text': generate_impact_text(selected_measures),
            'control_measures': generate_control_measures(selected_measures),
            'verification': generate_verification_text(selected_measures, has_high_risk),
            'has_high_risk': has_high_risk
        }
        
        doc = create_ats_document(data)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        new_balance = balance - 10.00
        set_user_credits(user_id, new_balance)
        
        clean_address = address.replace(',', '').replace(' ', '_')[:50]
        filename = f"ATS_{clean_address}_Annex_8_2_35.docx"
        
        return StreamingResponse(
            file_stream,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )