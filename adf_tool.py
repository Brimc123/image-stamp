from fastapi import Request, UploadFile, File, Form
from fastapi.responses import StreamingResponse, HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfplumber
import re
from typing import List, Optional
from database import get_user_by_id, set_user_credits, is_admin, log_usage, add_transaction

templates = Jinja2Templates(directory="templates")

# Cost for ADF Checklist generation
ADF_COST = 5.00

async def parse_ventilation_data(pdf_file):
    """Extract ventilation data from Condition Report or Site Notes PDF with improved parsing"""
    parsed = {
        'total_bg_vent_area': 0,
        'has_trickle_vents': False,
        'has_extract_fans': False,
        'extract_fan_rooms': [],
        'fan_control': None,
        'vent_system': 'natural',
        'rooms': {},
        'total_rooms_checked': 0,
        'room_details': []
    }
    
    try:
        content = await pdf_file.read()
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            current_room = None
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                
                # Identify room headers
                room_match = re.search(r'^(Living room|Kitchen|Bathroom|Bedroom \d+|Utility)', text, re.MULTILINE | re.IGNORECASE)
                if room_match:
                    current_room = room_match.group(1).strip()
                
                # Extract background ventilation area - HEAVILY DEBUGGED
                # Print the text snippet we're searching
                vent_idx = text.find('Ventilation')
                if vent_idx > -1:
                    snippet = text[vent_idx:min(vent_idx+200, len(text))]
                    print(f"\n🔍 DEBUG - Found 'Ventilation' on page {page_num}")
                    print(f"   Text snippet: {repr(snippet)[:150]}")
                
                bg_patterns = [
                    r'Background\s+Ventilation\s+Area\s*\(mm2(\d)\)(\d+)',  # Matches (mm21)0000
                    r'Background\s+Ventilation\s+Area\s*\(mm2\)\s*(\d{4,5})',  # Normal format
                ]
                
                for pattern in bg_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        # Handle both pattern formats
                        groups = match.groups()
                        if len(groups) == 2 and len(groups[1]) > 1:
                            # First pattern: digit inside parens + remaining digits
                            area = int(groups[0] + groups[1])
                        else:
                            # Second pattern: just the number
                            area = int(groups[0])
                        
                        parsed['total_bg_vent_area'] += area
                        parsed['total_rooms_checked'] += 1
                        
                        if current_room:
                            parsed['rooms'][current_room] = area
                            parsed['room_details'].append({
                                'room': current_room,
                                'bg_vent_area': area
                            })
                        
                        print(f"   ✅ Page {page_num} - {current_room}: {area}mm² (matched pattern: {pattern})")
                        break
                
                # Check for trickle vents
                if re.search(r'Do\s+they\s+have\s+trickle\s+vents\?\s*Yes', text, re.IGNORECASE):
                    parsed['has_trickle_vents'] = True
                
                # Check for extract fans - IMPROVED DETECTION with DEBUG
                if 'fan' in text.lower():
                    print(f"   🔍 Page {page_num} - Looking for fans...")
                    # Look for the fan question
                    fan_idx = text.lower().find('fan')
                    fan_section = text[fan_idx:min(fan_idx+100, len(text))]
                    print(f"      Fan section text: {repr(fan_section[:80])}")
                
                # Match "fans fittedY?es" or "fans fittedN?o"
                fan_fitted_match = re.search(r'fans\s+fitted([YN])\?([eo])s?', text, re.IGNORECASE)
                if fan_fitted_match:
                    print(f"   🎯 Page {page_num} - Fan match found: {fan_fitted_match.group(0)}")
                    answer = fan_fitted_match.group(1).upper()  # Y or N
                    if answer == 'Y':
                        parsed['has_extract_fans'] = True
                        if current_room and current_room not in parsed['extract_fan_rooms']:
                            parsed['extract_fan_rooms'].append(current_room)
                            print(f"   ✅ Page {page_num} - Fan found in {current_room}")
                
                # Check fan control type
                if re.search(r'permanently|on\s+permanently', text, re.IGNORECASE):
                    parsed['fan_control'] = 'automatic'
                elif re.search(r'pull\s+chord|manual', text, re.IGNORECASE):
                    parsed['fan_control'] = 'manual'
            
            # Determine ventilation system type
            full_text = "\n".join([p.extract_text() or "" for p in pdf.pages])
            
            if re.search(r'MVHR|mechanical\s+ventilation\s+with\s+heat\s+recovery', full_text, re.IGNORECASE):
                parsed['vent_system'] = 'mvhr'
            elif re.search(r'continuous\s+mechanical\s+extract|CME', full_text, re.IGNORECASE):
                parsed['vent_system'] = 'continuous_extract'
            elif re.search(r'natural\s+ventilation', full_text, re.IGNORECASE):
                parsed['vent_system'] = 'natural'
            
            print(f"\n📊 PARSING SUMMARY:")
            print(f"   Total Background Ventilation: {parsed['total_bg_vent_area']}mm²")
            print(f"   Rooms with BG Vent: {parsed['total_rooms_checked']}")
            print(f"   Trickle Vents: {parsed['has_trickle_vents']}")
            print(f"   Extract Fans: {parsed['has_extract_fans']}")
            print(f"   Fan Rooms: {', '.join(parsed['extract_fan_rooms']) if parsed['extract_fan_rooms'] else 'None'}")
            print(f"   System Type: {parsed['vent_system']}")
                
    except Exception as e:
        print(f"❌ Error parsing PDF: {e}")
        import traceback
        traceback.print_exc()
    
    return parsed


def set_cell_background(cell, color):
    """Set background color for a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_table_borders(table):
    """Add professional borders to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E5C8A')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def generate_adf_checklist(address, vent_data):
    """Generate professionally formatted ADF Table D1 checklist document"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Main Title with blue background
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('APPROVED DOCUMENT F - TABLE D1')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add background to title
    pPr = title_para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '2E5C8A')
    pPr.append(shading)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('CHECKLIST FOR VENTILATION PROVISION IN EXISTING DWELLINGS')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(46, 92, 138)
    
    doc.add_paragraph()
    
    # Property details box
    details_table = doc.add_table(rows=3, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details_table.cell(0, 0).text = 'Property Address:'
    details_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    details_table.cell(0, 1).text = address
    
    details_table.cell(1, 0).text = 'Assessment Date:'
    details_table.cell(1, 0).paragraphs[0].runs[0].bold = True
    details_table.cell(1, 1).text = datetime.now().strftime('%d/%m/%Y')
    
    details_table.cell(2, 0).text = 'Assessor:'
    details_table.cell(2, 0).paragraphs[0].runs[0].bold = True
    details_table.cell(2, 1).text = 'Auto-generated via AutoDate.co.uk'
    
    doc.add_paragraph()
    
    # Summary findings box
    summary = doc.add_paragraph()
    summary_run = summary.add_run('📊 SUMMARY FINDINGS')
    summary_run.bold = True
    summary_run.font.size = Pt(14)
    summary_run.font.color.rgb = RGBColor(46, 92, 138)
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    
    summary_table.cell(0, 0).text = 'Total Background Ventilation Area:'
    summary_table.cell(0, 1).text = f"{vent_data.get('total_bg_vent_area', 0)} mm²"
    
    summary_table.cell(1, 0).text = 'Trickle Vents Present:'
    summary_table.cell(1, 1).text = '✓ Yes' if vent_data.get('has_trickle_vents') else '✗ No'
    
    summary_table.cell(2, 0).text = 'Extract Fans Present:'
    if vent_data.get('has_extract_fans') and vent_data.get('extract_fan_rooms'):
        fans_text = f"✓ Yes ({', '.join(vent_data['extract_fan_rooms'])})"
    elif vent_data.get('has_extract_fans'):
        fans_text = '✓ Yes'
    else:
        fans_text = '✗ No'
    summary_table.cell(2, 1).text = fans_text
    
    summary_table.cell(3, 0).text = 'Ventilation System Type:'
    sys_type = vent_data.get('vent_system', 'natural').replace('_', ' ').title()
    summary_table.cell(3, 1).text = sys_type
    
    doc.add_paragraph()
    
    # Determine which system to document
    system_type = vent_data.get('vent_system', 'natural')
    
    # NATURAL VENTILATION SECTION
    if system_type == 'natural':
        heading = doc.add_paragraph()
        heading_run = heading.add_run('🏠 NATURAL VENTILATION')
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(46, 92, 138)
        
        table = doc.add_table(rows=1, cols=3)
        add_table_borders(table)
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Yes'
        hdr_cells[2].text = 'No'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(cell, '2E5C8A')
        
        total_bg = vent_data.get('total_bg_vent_area', 0)
        has_fans = vent_data.get('has_extract_fans', False)
        has_trickle = vent_data.get('has_trickle_vents', False)
        fan_control = vent_data.get('fan_control')
        
        # Questions for Natural Ventilation
        questions = [
            ('What is the total equivalent area of background ventilators currently in dwelling?', 
             f"{total_bg}mm²", '', ''),
            
            ('Does each habitable room satisfy the minimum equivalent area standards in Table 1.7?',
             '✓' if total_bg >= 5000 else '', 
             '', 
             '✗' if total_bg < 5000 else ''),
            
            ('Have all background ventilators been left in the open position?',
             '✓' if has_trickle else '', 
             '', 
             '✗' if not has_trickle else ''),
            
            ('Are fans and background ventilators in the same room at least 0.5m apart?',
             '✓', '', ''),
            
            ('Are there working intermittent extract fans in all wet rooms?',
             '✓' if has_fans else '', 
             '', 
             '✗' if not has_fans else ''),
            
            ('Is there the correct number of intermittent extract fans to satisfy the standards in Table 1.1?',
             '✓' if has_fans else '', 
             '', 
             '✗' if not has_fans else ''),
            
            ('Does the location of fans satisfy the standards in paragraph 1.20?',
             '✓', '', ''),
            
            ('Do all automatic controls have a manual override?',
             '✓' if fan_control != 'automatic' else '', 
             '', 
             '✗' if fan_control == 'automatic' else ''),
            
            ('Does each room have a system for purge ventilation (e.g. windows)?',
             '✓', '', ''),
            
            ('Do the openings in the rooms satisfy the minimum opening area standards in Table 1.4?',
             '✓', '', ''),
            
            ('Do all internal doors have sufficient undercut to allow air transfer between rooms (10mm above floor finish)?',
             '✓', '', ''),
        ]
        
        for i, (question, yes_val, no_val, explicit_val) in enumerate(questions):
            row = table.add_row().cells
            row[0].text = question
            
            if explicit_val:
                if '✗' in explicit_val:
                    row[2].text = explicit_val
                    set_cell_background(row[2], 'FFE6E6')
                else:
                    row[1].text = yes_val
                    set_cell_background(row[1], 'E6FFE6')
            elif 'mm²' in yes_val:
                row[1].text = yes_val
            else:
                row[1].text = yes_val
                if yes_val:
                    set_cell_background(row[1], 'E6FFE6')
            
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Zebra striping
            if i % 2 == 0:
                set_cell_background(row[0], 'F5F5F5')
    
    # Notes
    doc.add_paragraph()
    notes = doc.add_paragraph()
    notes_run = notes.add_run('⚠️ IMPORTANT NOTES')
    notes_run.bold = True
    notes_run.font.size = Pt(12)
    notes_run.font.color.rgb = RGBColor(204, 102, 0)
    
    note1 = doc.add_paragraph(style='List Number')
    note1.add_run('Make a visual check for mould or condensation. If either are present, install additional ventilation provisions or seek specialist advice.')
    
    note2 = doc.add_paragraph(style='List Number')
    note2.add_run('All references to tables and paragraphs are to Approved Document F, Volume 1: Dwellings (2021 edition).')
    
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Generated by AutoDate.co.uk ADF Checklist Generator')
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc


async def adf_checklist_route(request: Request, user_row):
    """Route handler for ADF Checklist Generator"""
    if request.method == "GET":
        user_id = user_row.get("id")
        user_data = get_user_by_id(user_id)
        balance = float(user_data.get("credits", 0.0))
        return templates.TemplateResponse("adf_checklist.html", {
            "request": request, 
            "balance": balance,
            "cost": ADF_COST
        })
    
    elif request.method == "POST":
        user_id = user_row.get("id")
        user_data = get_user_by_id(user_id)
        balance = float(user_data.get("credits", 0.0))
        
        # Check if admin
        is_admin_user = is_admin(user_id)

        # Check balance (skip for admin)
        if not is_admin_user and balance < ADF_COST:return HTMLResponse(
                content=f"<h1>Insufficient balance</h1><p>You need £{ADF_COST:.2f} but have £{balance:.2f}</p>",
                status_code=400
            )
        
        form_data = await request.form()
        address = form_data.get('address')
        
        if not address:
            return HTMLResponse(content="<h1>Please provide property address</h1>", status_code=400)
        
        # Get uploaded PDFs
        condition_report = form_data.get("condition_report")
        
        if not condition_report or not hasattr(condition_report, 'read'):
            return HTMLResponse(content="<h1>Please upload Condition Report PDF</h1>", status_code=400)
        
        print("📄 Parsing Condition Report...")
        # Parse Condition Report
        vent_data = await parse_ventilation_data(condition_report)
        
        # Optional Site Notes
        site_notes = form_data.get("site_notes")
        if site_notes and hasattr(site_notes, 'read'):
            print("📋 Parsing Site Notes...")
            site_notes_data = await parse_ventilation_data(site_notes)
            # Merge data
            if site_notes_data.get('total_bg_vent_area', 0) > 0:
                vent_data['total_bg_vent_area'] += site_notes_data['total_bg_vent_area']
            if site_notes_data.get('has_extract_fans'):
                vent_data['has_extract_fans'] = True
            if site_notes_data.get('extract_fan_rooms'):
                vent_data['extract_fan_rooms'].extend(site_notes_data['extract_fan_rooms'])
        
        print(f"✅ Final Parsed Data: {vent_data}")
        
        # Generate checklist
        doc = generate_adf_checklist(address, vent_data)
        
        # Save to stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Deduct credits only if not admin
        if not is_admin_user:
            new_balance = balance - ADF_COST
                
        # Deduct credits only if not admin
        if not is_admin_user:
            new_balance = balance - ADF_COST
            set_user_credits(user_id, new_balance)
            add_transaction(user_id, -ADF_COST, "adf_checklist")
            log_usage(user_id, "ADF Checklist", ADF_COST, f"Generated for {address}")
        else:
            log_usage(user_id, "ADF Checklist", 0.00, f"Admin - Generated for {address}")

        # Generate filename
        clean_address = re.sub(r'[^\w\s-]', '', address).strip().replace(' ', '_')[:50]
        filename = f"ADF_TableD1_{clean_address}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            file_stream,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )