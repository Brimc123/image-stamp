"""
PAS 2035 Compliance Documents Generator
RC Consultants - Professional Retrofit Documentation
ULTRA PREMIUM EDITION - Maximum visual impact and comprehensive content
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io

# Measure descriptions
MEASURE_DESCRIPTIONS = {
    "RIR": "Roof Insulation and Repair",
    "IWI": "Internal Wall Insulation",
    "LOFT": "Loft Insulation",
    "GAS_BOILER": "Gas Boiler Replacement",
    "ESH": "Electric Storage Heaters",
    "ASHP": "Air Source Heat Pump",
    "SOLAR": "Solar PV Panels",
    "CWI": "Cavity Wall Insulation",
    "UFI": "Underfloor Insulation",
    "HEATING_CONTROLS": "Heating Controls Upgrade"
}

# COMPREHENSIVE measure descriptions (500+ words each)
MEASURE_DETAILS = {
    "RIR": """
ROOF INSULATION AND REPAIR - COMPREHENSIVE OVERVIEW

Your roof insulation and repair works represent a critical investment in your home's thermal performance and structural integrity. Heat rises, and in an uninsulated home, up to 25% of all heat generated can escape through the roof space. This substantial heat loss not only increases your energy bills but also contributes significantly to your carbon footprint.

TECHNICAL SPECIFICATIONS:
The installation includes high-performance insulation materials exceeding current Building Regulations Part L requirements (minimum 270mm for cold roofs, 100-150mm for warm roofs depending on material type). We use only certified materials with independently verified thermal performance values, ensuring long-term reliability and effectiveness.

INSTALLATION PROCESS:
Our certified installers follow PAS 2030:2023 standards throughout the installation. This includes:
- Comprehensive roof space inspection and preparation
- Identification and protection of all electrical cables and penetrations
- Installation of appropriate ventilation to prevent condensation
- Careful placement of insulation to achieve consistent coverage with no gaps or cold bridges
- Sealing of all penetrations and potential air leakage points
- Installation of loft hatches with insulated covers where appropriate
- Clear marking of safe access routes and any restrictions

THERMAL PERFORMANCE:
The improved U-value (measure of heat loss) typically reduces from 2.3 W/m²K (uninsulated) to 0.16 W/m²K or better, representing a dramatic improvement in thermal efficiency. This translates to:
- Reduction in heating energy consumption of 15-25% for a typical property
- Improved thermal comfort with more stable indoor temperatures
- Reduced risk of condensation and associated issues
- Better temperature regulation during summer months (reduced overheating)

ADDITIONAL BENEFITS:
Beyond energy savings, roof insulation provides noise reduction from external sources (rain, aircraft, traffic), improved fire resistance (most modern insulation materials are fire-rated), increased property value, and eligibility for improved Energy Performance Certificate (EPC) ratings.

MAINTENANCE AND LONGEVITY:
Quality roof insulation typically lasts 40+ years with minimal maintenance required. We recommend periodic visual inspections (every 5 years) to check for any disturbance, water ingress, or pest activity. The insulation should remain dry and undisturbed for optimal performance throughout its lifespan.

ENVIRONMENTAL IMPACT:
Your roof insulation will prevent approximately 0.5-1.0 tonnes of CO2 emissions annually, equivalent to taking a car off the road for 2,000-4,000 miles each year. Over the 40-year lifespan, this represents a substantial contribution to carbon reduction targets while saving you thousands of pounds in energy costs.
""",

    "IWI": """
INTERNAL WALL INSULATION - DETAILED TECHNICAL GUIDE

Internal wall insulation (IWI) represents one of the most effective retrofit solutions for properties where external wall insulation is not suitable due to planning constraints, property type, or aesthetic considerations. Solid wall properties, which make up approximately 7 million homes in the UK, typically lose 33% of their heat through uninsulated walls - making wall insulation the single most impactful energy efficiency measure available.

SYSTEM SPECIFICATION:
Your internal wall insulation system has been carefully selected based on your property's specific characteristics. The installation typically includes:
- Insulated plasterboard (commonly 52.5-72.5mm total thickness) or alternative insulation board systems
- Integrated or separate vapour control layer to manage moisture movement
- Mechanical fixing and/or adhesive bonding as appropriate for your wall type
- Professional finishing with skim plaster coat ready for decoration
- Careful detailing around all penetrations, corners, and junctions
- Consideration of thermal bridging at reveals and junctions

INSTALLATION METHODOLOGY:
The installation follows a precise sequence developed to minimize disruption while ensuring optimal performance:

1. PREPARATION PHASE:
All furniture and belongings are protected or temporarily relocated. Existing fixtures (radiators, sockets, switches) are carefully removed and their positions recorded. The existing wall surface is prepared to ensure good adhesion.

2. INSTALLATION PHASE:
Insulation boards are cut precisely and fixed to walls following manufacturer specifications. Special attention is paid to corners, reveals, and junctions to eliminate thermal bridges. All penetrations for electrical outlets, switches, and services are carefully cut and sealed. Vapour control layers are installed with all joints properly sealed.

3. FINISHING PHASE:
Professional plastering provides a smooth, high-quality finish ready for decoration. All fixtures are reinstalled in their original positions (or relocated as agreed). Final inspection ensures all details meet PAS 2030:2023 requirements.

THERMAL PERFORMANCE IMPROVEMENT:
The U-value improvement is substantial, typically from 2.1 W/m²K (uninsulated solid wall) to 0.30 W/m²K or better with IWI. This represents:
- 60-70% reduction in heat loss through walls
- Elimination of cold spots and condensation risk on internal wall surfaces
- Warmer wall surface temperatures (typically 3-5°C warmer)
- More even temperature distribution throughout rooms
- Reduced heating demand of 15-35% depending on property type

CONSIDERATIONS AND ADAPTATIONS:
IWI reduces internal room dimensions by approximately 80-100mm per insulated wall. We provide careful planning to minimize impact:
- Electrical outlets and switches are brought forward to new wall surface
- Skirting boards are reinstalled or replaced
- Window reveals are insulated and finished professionally
- Door frames may require extension pieces
- Radiator brackets are extended where necessary

MOISTURE MANAGEMENT:
Modern IWI systems include sophisticated moisture management strategies. The vapour control layer prevents internal moisture from reaching the cold wall surface while allowing any moisture in the wall to dry outward. This approach, combined with adequate ventilation, ensures long-term performance without moisture-related issues.

ACOUSTIC BENEFITS:
As a valuable secondary benefit, IWI significantly improves sound insulation. The additional mass and insulation layer can reduce sound transmission through external walls by 40-50 decibels, creating a quieter, more peaceful internal environment.

COMPATIBILITY WITH ORIGINAL FEATURES:
In properties with original features, we take extra care to preserve character while improving performance. This includes careful detailing around decorative moldings, sensitive handling of historic plasterwork, and consultation regarding the impact on original features.

LONG-TERM PERFORMANCE:
With proper installation and adequate ventilation, IWI systems provide excellent performance for 50+ years. The materials are stable, non-degrading, and require no maintenance beyond normal decoration cycles. Regular ventilation (particularly in bathrooms and kitchens) ensures optimal performance and prevents moisture issues.

FINANCIAL BENEFITS:
Beyond immediate energy bill savings (typically £200-400 annually for a typical property), IWI improves property value, increases EPC rating (often by 2-3 bands), makes homes more attractive to buyers, and provides long-term protection against rising energy costs.
""",

    "LOFT": """
LOFT INSULATION - COMPLETE TECHNICAL DOCUMENTATION

Loft insulation stands as the most cost-effective energy efficiency measure available to homeowners. With heat naturally rising, an uninsulated loft allows approximately 25% of your home's heat to escape directly through the roof - representing a substantial waste of energy and money. Modern loft insulation standards require depths of 270mm or more, significantly exceeding older standards and providing exceptional thermal performance.

MATERIAL SELECTION AND SPECIFICATION:
Your loft insulation uses premium materials selected for optimal performance:
- Mineral wool (glass or stone) meeting British Standard BS EN 13162
- Thermal conductivity (λ-value) of 0.035-0.044 W/mK
- Fire classification A1 or A2 (non-combustible)
- Water repellent treatment to resist moisture
- VOC-free and suitable for occupied dwellings
- 50-year minimum design life

COMPREHENSIVE INSTALLATION PROCESS:

STAGE 1 - PRE-INSTALLATION SURVEY AND PREPARATION:
Before any insulation is installed, our certified assessors conduct a thorough survey:
- Identification of all services (electrical cables, pipes, junction boxes)
- Assessment of ventilation requirements and existing ventilation paths
- Checking for any existing moisture or structural issues
- Measuring loft dimensions and calculating material quantities
- Planning access routes and storage requirements
- Identifying any areas requiring special attention (chimneys, hatches, eaves)

STAGE 2 - LOFT SPACE PREPARATION:
The loft is carefully prepared to ensure optimal insulation performance:
- All debris and old insulation material is removed if necessary
- Electrical cables are identified and marked for protection
- Any defective roofing is reported and remedied
- Water tanks and pipes are insulated separately if in cold loft space
- Adequate boarding is installed for access routes if required
- Eaves ventilation is checked and improved if necessary

STAGE 3 - INSULATION INSTALLATION:
The installation follows a precise methodology:
- First layer (if required) installed between ceiling joists
- Second layer installed across joists to achieve required depth
- All areas covered uniformly with no gaps or compression
- Electrical cables laid over insulation to prevent overheating
- Loft hatch insulated with dedicated insulation and draught seal
- Eaves ventilation maintained with appropriate ventilation gaps
- Any recessed light fittings protected with fire-rated covers
- Clear access routes maintained and marked

THERMAL PERFORMANCE AND ENERGY SAVINGS:
The improvement in thermal performance is dramatic and quantifiable:
- U-value improves from 2.3 W/m²K (uninsulated) to 0.16 W/m²K (270mm insulation)
- Heat loss through roof reduced by 93%
- Typical energy bill savings of £200-300 per year for average property
- Payback period of 2-5 years depending on installation cost
- Carbon emissions reduced by 0.5-1.0 tonnes CO2 per year

VENTILATION MANAGEMENT:
Adequate ventilation is critical for loft performance and building health:
- Eaves ventilation maintained with 25mm continuous gap
- Ridge ventilation checked and improved if necessary
- Gable end vents assessed and upgraded where required
- Cold water tanks moved to warm side of insulation if possible
- Ventilation paths kept clear to allow air circulation
- Moisture management strategy implemented

SPECIAL CONSIDERATIONS AND DETAILS:

Chimneys and Flues:
All chimneys and flues require special attention with insulation kept back by minimum 100mm to prevent fire risk. We use fire-resistant barriers where necessary and ensure adequate ventilation around any heat-producing services.

Recessed Lighting:
Modern LED fittings generate less heat but still require attention. We install fire-rated covers (rated for minimum 30 minutes fire resistance) over any recessed lights, maintaining minimum air gaps as specified by manufacturers.

Hatches and Access Points:
Loft hatches represent significant heat loss points if not properly insulated. We fit:
- 100mm+ insulation to hatch cover
- Draught seal around hatch frame
- Secure catches to maintain seal when closed
- Easy-open mechanisms for convenient access

LONG-TERM PERFORMANCE AND MAINTENANCE:
Quality loft insulation requires virtually no maintenance and performs effectively for 40+ years. We recommend:
- Visual inspection every 5 years to check for disturbance
- Checking for any water ingress or moisture issues
- Ensuring ventilation paths remain clear
- Not compressing insulation when storing items in loft
- Maintaining safe access routes

ADDITIONAL BENEFITS BEYOND ENERGY SAVING:
- Summer cooling: Loft insulation reduces heat gain during summer months
- Noise reduction: Significant reduction in external noise (rain, aircraft, traffic)
- Property value: Improved EPC rating increases property value and marketability
- Comfort: More stable temperatures throughout home
- Condensation control: Warmer ceiling surface reduces condensation risk

ENVIRONMENTAL IMPACT AND SUSTAINABILITY:
Your loft insulation contributes significantly to environmental protection:
- Prevents 0.7 tonnes CO2 emissions annually
- Equivalent to 28,000 miles of carbon offset over 40-year life
- Supports UK net-zero carbon targets
- Made from recycled materials (most mineral wool contains 80%+ recycled content)
- Fully recyclable at end of life

COMPLIANCE AND CERTIFICATION:
Installation completed in accordance with:
- Building Regulations Part L (Conservation of Fuel and Power)
- PAS 2030:2023 (Installer competency)
- PAS 2035:2023 (Retrofit coordination)
- British Standards BS 5250 (Moisture control)
- Manufacturer installation guidelines
- TrustMark quality framework
""",

    "GAS_BOILER": """
GAS BOILER REPLACEMENT - COMPLETE INSTALLATION GUIDE

Your new A-rated condensing gas boiler represents modern heating technology at its finest. Replacing an old, inefficient boiler (typically 60-70% efficient or worse) with a modern condensing boiler (minimum 90% efficient, often 94%+ in real-world conditions) is one of the most impactful home improvements available. Over the 15-20 year lifespan of your new boiler, you'll save thousands of pounds in energy costs while enjoying significantly improved comfort, reliability, and control.

BOILER SPECIFICATION AND TECHNOLOGY:

Your New Boiler Features:
- A-rated condensing technology achieving 92-94% efficiency
- Modulating burner that adjusts output to match demand precisely
- Weather compensation capability (if external sensor fitted)
- Integrated frost protection
- Built-in pump and expansion vessel
- Digital display showing status and diagnostics
- Quiet operation (typically <40dB)
- Compact design occupying less space than older models

Technical Specifications:
- Output: [Sized specifically for your property's heat loss calculation]
- Fuel: Natural gas (compliant with Gas Safety Regulations 1998)
- Flue type: Room sealed condensing
- Controls compatible: OpenTherm, on/off, weather compensation
- DHW performance: Combi models provide instant hot water; system boilers work with cylinder
- Warranty: [Typically 5-10 years depending on model and registration]

COMPREHENSIVE INSTALLATION PROCESS:

STAGE 1 - PRE-INSTALLATION ASSESSMENT:
Before installation commences, our Gas Safe registered engineers conduct thorough assessment:
- Accurate heat loss calculation determining correct boiler size
- Gas supply adequacy check (ensuring sufficient pressure and flow rate)
- Flue routing options assessed and optimal position determined
- Condensate disposal arrangements planned
- Control strategy discussed and agreed
- Hot water requirements confirmed (cylinder size, flow rates)
- Existing system condition assessed

STAGE 2 - SYSTEM PREPARATION:
Proper preparation ensures long boiler life and optimal performance:
- Power flush of existing heating system to remove debris and magnetite
- Chemical cleansing to remove scale and deposits
- Magnetic filter installation to protect new boiler from debris
- Chemical inhibitor added to system water to prevent corrosion
- All radiators checked and bled
- Existing valves assessed and replaced if worn
- Radiator valves upgraded to thermostatic (TRVs) where beneficial

STAGE 3 - BOILER INSTALLATION:
Installation follows strict Gas Safe and manufacturer requirements:
- Old boiler safely disconnected and removed
- New boiler mounting bracket securely fixed to wall
- Gas supply pipework checked, tested, and connected
- Central heating flow and return pipes connected
- Pressure relief valve and condensate drain properly installed
- Electrical connection made by qualified electrician (fused spur)
- Flue correctly positioned with terminal guard where required
- System filled, pressurized, and thoroughly checked for leaks

STAGE 4 - COMMISSIONING AND TESTING:
Comprehensive commissioning ensures safe, efficient operation:
- Gas pressure checked at appliance and adjusted if necessary
- Combustion analysis performed and recorded
- Burner pressure set and verified
- Safety devices tested (overheat stat, pressure relief, flame failure)
- Full system operational test over complete heating cycle
- Hot water delivery tested (combi) or cylinder heat-up tested (system)
- Controls commissioned and demonstrated to householder
- All documentation completed including Benchmark checklist

CONTROL SYSTEM AND OPTIMAL OPERATION:

Modern Controls for Maximum Efficiency:
Your new boiler works with advanced controls to minimize energy consumption:
- Programmable room thermostat allowing multiple time/temperature settings
- Thermostatic radiator valves (TRVs) on individual radiators
- Weather compensation (if fitted) adjusting flow temperature based on outside temperature
- OpenTherm digital communication for smooth modulation
- Frost protection maintaining minimum temperature when property unoccupied

Optimal Operating Strategy:
To achieve maximum efficiency and comfort:
- Set room thermostat to comfortable temperature (typically 19-21°C)
- Use TRVs to reduce temperatures in lesser-used rooms
- Utilize timer to heat home only when occupied
- Avoid "boost" override unless absolutely necessary
- Reduce night-time temperature by 2-3°C for sleeping comfort and energy saving
- Consider "set-back" temperature for extended absences (12-15°C)

EFFICIENCY AND PERFORMANCE:

Real-World Efficiency:
Modern condensing boilers achieve their high efficiency through:
- Condensing technology extracting latent heat from flue gases
- Modulating burner operating at lowest output necessary
- Well-insulated heat exchanger minimizing heat loss
- Efficient pump with variable speed operation
- Reduced standing losses due to better insulation

Typical Performance Figures:
- Seasonal efficiency: 88-92% (accounting for cycling losses)
- Part-load efficiency: Often exceeds full-load efficiency due to condensing operation
- DHW efficiency: 85-90% (combi boilers)
- Parasitic electrical consumption: <100W (pump, controls, display)

ENERGY SAVINGS AND FINANCIAL BENEFITS:

Quantified Savings:
Replacing an old boiler (G-rated, 60% efficiency) with new A-rated boiler (92% efficiency):
- Energy consumption reduced by approximately 35%
- Typical annual saving: £300-500 depending on property size and usage
- Payback period: 5-8 years from energy savings alone
- Carbon emissions reduced by 1.0-1.5 tonnes CO2 annually
- Over 15-year boiler life: Total savings of £4,500-7,500

Additional Financial Benefits:
- Reduced repair and maintenance costs compared to old boiler
- Improved property value and marketability
- Better EPC rating (typically improving by 1-2 bands)
- Eligibility for various government schemes and incentives
- Lower insurance premiums (some insurers offer discounts)

MAINTENANCE AND SERVICE REQUIREMENTS:

Annual Service Essentials:
To maintain warranty and ensure safe, efficient operation:
- Annual service by Gas Safe registered engineer (typically £80-120)
- Combustion analysis and adjustment
- Safety device testing
- Condensate trap cleaning
- Pressure check and repressurise if needed
- Visual inspection of flue and terminal
- Controls operation check
- Benchmark logbook updated

Homeowner Maintenance:
Simple tasks to maintain performance:
- Check pressure gauge monthly (should be 1.0-1.5 bar when cold)
- Bleed radiators if cold spots develop
- Keep boiler area clear and well-ventilated
- Check condensate pipe in winter (ensure not frozen)
- Test controls periodically
- Report any unusual noises, smells, or error codes immediately

SAFETY FEATURES AND PROTECTION:

Multiple Safety Systems:
Your boiler incorporates numerous safety features:
- Flame failure device shuts off gas if flame extinguished
- Overheat thermostat prevents excessive temperature
- Pressure relief valve protects against overpressure
- Flue gas spillage detection (if fitted)
- Frost protection prevents system freezing
- Dry fire protection (combi models)

Carbon Monoxide Protection:
For complete peace of mind:
- Room sealed design eliminates risk of flue gas spillage
- Interlock systems prevent operation if unsafe
- We strongly recommend carbon monoxide detector installation
- Annual service includes combustion analysis

ENVIRONMENTAL IMPACT:

Carbon Reduction:
Your new boiler significantly reduces carbon emissions:
- Approximately 1.2 tonnes CO2 saved annually
- Over 15-year life: 18 tonnes CO2 prevented
- Equivalent to taking car off road for 50,000+ miles
- Supports UK legally binding net-zero targets

Future Considerations:
Modern condensing boilers can potentially operate on hydrogen blends:
- Most new boilers are "hydrogen-ready" awaiting infrastructure
- Firmware updates may enable hydrogen operation in future
- Positioned for smooth transition to low-carbon heating
""",

    # Continue with other measures...
    "ESH": """[Similar comprehensive 500+ word content for Electric Storage Heaters]""",
    "ASHP": """[Similar comprehensive 500+ word content for Air Source Heat Pump]""",
    "SOLAR": """[Similar comprehensive 500+ word content for Solar PV]""",
    "CWI": """[Similar comprehensive 500+ word content for Cavity Wall Insulation]""",
    "UFI": """[Similar comprehensive 500+ word content for Underfloor Insulation]""",
    "HEATING_CONTROLS": """[Similar comprehensive 500+ word content for Heating Controls]"""
}


def set_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_decorative_border(paragraph):
    """Add decorative border to paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '0066CC')
        pBdr.append(border)
    
    pPr.append(pBdr)


def generate_pas2035_documents(form_data: dict):
    """Generate all 3 documents"""
    rc_id = form_data.get('rc_id', '')
    rc_name = form_data.get('rc_name', '')
    property_address = form_data.get('property_address', '')
    customer_name = form_data.get('customer_name', '')
    customer_address = form_data.get('customer_address', '')
    installer_name = form_data.get('installer_name', '')
    installer_contact = form_data.get('installer_contact', '')
    measures = form_data.get('measures', [])
    project_date = form_data.get('project_date', datetime.now().strftime('%d/%m/%Y'))
    install_start_date = form_data.get('install_start_date', '')
    conflict_of_interest = form_data.get('conflict_of_interest', 'No')
    
    measure_names = [MEASURE_DESCRIPTIONS.get(m, m) for m in measures]
    measures_text = ", ".join(measure_names)
    
    sf48_doc = generate_sf48_certificate(rc_id, rc_name, property_address, measures_text, project_date)
    intro_doc = generate_intro_letter(customer_name, customer_address, measures_text, install_start_date, installer_contact)
    handover_doc = generate_handover_letter(customer_name, customer_address, measures_text, project_date, rc_name, installer_name, conflict_of_interest, measures)
    
    return sf48_doc, intro_doc, handover_doc

def add_document_border(section):
    """Add border around entire document"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '0066CC')
        pgBorders.append(border)
    
    sectPr.append(pgBorders)


def generate_sf48_certificate(rc_id, rc_name, property_address, measures_text, project_date):
    """Generate ONE PAGE SF48 Certificate - clean and professional"""
    doc = Document()
    
    # Add document border
    add_document_border(doc.sections[0])
    
    # HEADER
    header_para = doc.add_paragraph()
    header_run = header_para.add_run('RC CONSULTANTS')
    header_run.font.size = Pt(14)
    header_run.bold = True
    header_run.font.color.rgb = RGBColor(0, 51, 153)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tagline = doc.add_paragraph('Professional Retrofit Coordination Services')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].font.size = Pt(9)
    tagline.runs[0].italic = True
    tagline.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    # Contact line
    contact = doc.add_paragraph('202 Queens Dock Business Centre, Liverpool, L1 0BG | ☎ 0800 001 6127')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(8)
    contact.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # TITLE
    title = doc.add_paragraph()
    title_run = title.add_run('RETROFIT PROJECT\nCLAIM OF COMPLIANCE')
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(204, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Based on Self-Assessment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # Compliance statement
    compliance = doc.add_paragraph()
    compliance.add_run('This Retrofit Project undertaken at the below address was completed in accordance with ')
    compliance.add_run('PAS 2035:2023').bold = True
    compliance.runs[1].font.color.rgb = RGBColor(0, 102, 204)
    compliance.add_run(' and ')
    compliance.add_run('PAS 2030:2023').bold = True
    compliance.runs[3].font.color.rgb = RGBColor(0, 102, 204)
    compliance.add_run(' by the Retrofit Coordinator identified below.')
    for run in compliance.runs:
        run.font.size = Pt(10)
    compliance.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # PROJECT DETAILS TABLE
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)
    
    rows_data = [
        ('Retrofit Coordinator ID:', rc_id),
        ('Retrofit Coordinator Name:', rc_name),
        ('Property Address:', property_address),
        ('Measures Installed:', measures_text),
        ('Completion Date:', project_date),
        ('Coordinator Signature:', '')
    ]
    
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        
        label_cell = row.cells[0]
        label_cell.text = label
        label_para = label_cell.paragraphs[0]
        label_para.runs[0].font.bold = True
        label_para.runs[0].font.size = Pt(10)
        set_cell_background(label_cell, 'E7F2FA')
        
        value_cell = row.cells[1]
        value_cell.text = value if value else ''
        value_para = value_cell.paragraphs[0]
        value_para.runs[0].font.size = Pt(10)
        
        if i == 5:
            row.height = Inches(0.8)
    
    doc.add_paragraph()
    
    # CERTIFICATION
    cert_heading = doc.add_paragraph('CERTIFICATION STATEMENT')
    cert_heading.runs[0].bold = True
    cert_heading.runs[0].font.size = Pt(11)
    cert_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    cert_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cert_text = doc.add_paragraph()
    cert_text.add_run('I certify that the retrofit measures listed above have been completed in full compliance with PAS 2035:2023 and PAS 2030:2023 standards. All work has been carried out by appropriately qualified installers, and complete project documentation has been maintained in accordance with the relevant standards.')
    cert_text.runs[0].font.size = Pt(9)
    cert_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Standards compliance
    standards = doc.add_paragraph('STANDARDS COMPLIANCE')
    standards.runs[0].bold = True
    standards.runs[0].font.size = Pt(10)
    standards.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    standards_list = [
        '• PAS 2035:2023 - Retrofitting dwellings for improved energy efficiency',
        '• PAS 2030:2023 - Installation competency standards',
        '• Building Regulations Part L - Conservation of Fuel and Power',
        '• TrustMark Quality Framework'
    ]
    
    for item in standards_list:
        p = doc.add_paragraph(item)
        p.runs[0].font.size = Pt(8)
        p.left_indent = Inches(0.3)
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run('This document confirms compliance with PAS 2035 requirements for domestic retrofit projects.')
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].italic = True
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def generate_intro_letter(customer_name, customer_address, measures_text, install_start_date, installer_contact):
    """Generate clean introduction letter - NO installation instructions"""
    doc = Document()
    
    # Add document border
    add_document_border(doc.sections[0])
    
    # HEADER
    header = doc.add_paragraph()
    header_run = header.add_run('RC CONSULTANTS')
    header_run.bold = True
    header_run.font.size = Pt(14)
    header_run.font.color.rgb = RGBColor(0, 51, 153)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tagline = doc.add_paragraph('Professional Retrofit Coordination & Energy Consultancy')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].font.size = Pt(9)
    tagline.runs[0].italic = True
    tagline.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # Contact details
    contact_lines = [
        '202 Queens Dock Business Centre',
        'Norfolk House, Liverpool, L1 0BG',
        'Phone: 0800 001 6127',
        'Email: info@rcconsultants.co.uk'
    ]
    
    for line in contact_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Customer address
    for line in [customer_name] + (customer_address.split('\n') if customer_address else []):
        p = doc.add_paragraph(line)
        if p.runs:
            p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph(datetime.now().strftime('%d %B %Y'))
    date_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Greeting
    greeting = doc.add_paragraph(f'Dear {customer_name},')
    greeting.runs[0].font.size = Pt(10)
    greeting.runs[0].bold = True
    
    doc.add_paragraph()
    
    # Subject
    subject = doc.add_paragraph()
    subject.add_run('RE: Retrofit Installation Project - ')
    subject.add_run(measures_text).bold = True
    for run in subject.runs:
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Body
    p1 = doc.add_paragraph()
    p1.add_run('We are pleased to confirm that we will shortly begin the installation of ')
    p1.add_run(measures_text).bold = True
    p1.add_run(' at your property. This work will be carried out in full compliance with PAS 2030:2023 standards, coordinated under PAS 2035:2023 requirements.')
    for run in p1.runs:
        run.font.size = Pt(10)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    p2 = doc.add_paragraph()
    p2.add_run('Installation is scheduled to commence on ')
    p2.add_run(install_start_date).bold = True
    p2.add_run('. Our certified installers will work efficiently to minimize disruption to your daily routine while maintaining the highest standards of workmanship.')
    for run in p2.runs:
        run.font.size = Pt(10)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.add_run('If you have any questions or concerns regarding this project, please contact ')
    p3.add_run(installer_contact).bold = True
    p3.add_run(' or our office on 0800 001 6127.')
    for run in p3.runs:
        run.font.size = Pt(10)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    p4 = doc.add_paragraph()
    p4.add_run('We look forward to transforming your home with these energy-efficient improvements.')
    p4.runs[0].font.size = Pt(10)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Closing
    closing = doc.add_paragraph('Yours sincerely,')
    closing.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature = doc.add_paragraph('RC CONSULTANTS')
    signature.runs[0].font.size = Pt(11)
    signature.runs[0].bold = True
    signature.runs[0].font.color.rgb = RGBColor(0, 51, 153)
    
    team = doc.add_paragraph('Retrofit Coordination Team')
    team.runs[0].font.size = Pt(9)
    team.runs[0].italic = True
    
    return doc


def generate_handover_letter(customer_name, customer_address, measures_text, project_date, rc_name, installer_name, conflict_of_interest, measures):
    """Generate clean, professional handover document with enhanced borders"""
    doc = Document()
    
    # Set document font to Calibri
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Document border on first page
    add_document_border(doc.sections[0])
    
    # TITLE
    title = doc.add_paragraph()
    title_run = title.add_run('RETROFIT PROJECT HANDOVER')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 112, 192)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('PAS 2035:2023 Compliant Project')
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(20)
    
    # PROJECT INFORMATION TABLE
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Medium Grid 1 Accent 1'
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)
    
    table_data = [
        ('Retrofit Coordinator ID', 'TMLN 2826127'),
        ('Retrofit Coordinator Name', 'Brian Mckevitt MCIOB'),
        ('Retrofit Coordinator Contact', 'brian@rcconsultants.co.uk | 0800 001 6127'),
        ('Customer Name', customer_name if customer_name else ''),
        ('Property Address', customer_address if customer_address else ''),
        ('Funding Stream', 'ECO4'),
        ('Measures Installed', measures_text if measures_text else ''),
        ('Completion Date', project_date if project_date else ''),
        ('Installation Company', installer_name if installer_name else ''),
        ('Conflict of Interest', conflict_of_interest if conflict_of_interest else 'No')
    ]
    
    for i, (label, value) in enumerate(table_data):
        row = table.rows[i]
        
        label_cell = row.cells[0]
        label_cell.text = label
        label_para = label_cell.paragraphs[0]
        label_para.runs[0].font.bold = True
        label_para.runs[0].font.size = Pt(10)
        label_para.runs[0].font.name = 'Calibri'
        set_cell_background(label_cell, 'D9E2F3')
        
        value_cell = row.cells[1]
        value_cell.text = value
        value_para = value_cell.paragraphs[0]
        value_para.runs[0].font.size = Pt(10)
        value_para.runs[0].font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # COMPLETION STATEMENT
    completion_heading = doc.add_paragraph()
    completion_heading.add_run('Installation Completion Statement')
    completion_heading.runs[0].font.size = Pt(14)
    completion_heading.runs[0].font.bold = True
    completion_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    completion_heading.runs[0].font.name = 'Calibri'
    completion_heading.space_before = Pt(12)
    
    completion_text = doc.add_paragraph()
    completion_text.add_run('I am pleased to confirm that all retrofit measures have been successfully installed at your property. All work has been completed in full compliance with PAS 2035:2023 and PAS 2030:2023 standards by certified installers.')
    completion_text.runs[0].font.size = Pt(10)
    completion_text.runs[0].font.name = 'Calibri'
    completion_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # INSTALLED MEASURES
    measures_heading = doc.add_paragraph()
    measures_heading.add_run('Installed Measures')
    measures_heading.runs[0].font.size = Pt(14)
    measures_heading.runs[0].font.bold = True
    measures_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    measures_heading.runs[0].font.name = 'Calibri'
    
    for measure_code in measures:
        measure_name = MEASURE_DESCRIPTIONS.get(measure_code, measure_code)
        measure_para = doc.add_paragraph()
        measure_para.add_run(f'✓ {measure_name}')
        measure_para.runs[0].font.size = Pt(11)
        measure_para.runs[0].font.name = 'Calibri'
        measure_para.runs[0].font.color.rgb = RGBColor(0, 102, 0)
        measure_para.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # DOCUMENTATION
    docs_heading = doc.add_paragraph()
    docs_heading.add_run('Documentation Provided')
    docs_heading.runs[0].font.size = Pt(14)
    docs_heading.runs[0].font.bold = True
    docs_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    docs_heading.runs[0].font.name = 'Calibri'
    
    docs_list = [
        'User guides and operation manuals',
        'Manufacturer warranties and guarantees',
        'Installation certificates',
        'Insurance-backed guarantees (where applicable)',
        'Building Control certification (where required)'
    ]
    
    for item in docs_list:
        doc_para = doc.add_paragraph()
        doc_para.add_run(f'• {item}')
        doc_para.runs[0].font.size = Pt(10)
        doc_para.runs[0].font.name = 'Calibri'
        doc_para.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # MAINTENANCE GUIDANCE
    maint_heading = doc.add_paragraph()
    maint_heading.add_run('Maintenance Guidance')
    maint_heading.runs[0].font.size = Pt(14)
    maint_heading.runs[0].font.bold = True
    maint_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    maint_heading.runs[0].font.name = 'Calibri'
    
    maint_intro = doc.add_paragraph()
    maint_intro.add_run('To ensure optimal performance and maintain warranty validity, please follow these maintenance recommendations:')
    maint_intro.runs[0].font.size = Pt(10)
    maint_intro.runs[0].font.name = 'Calibri'
    maint_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    for measure_code in measures:
        measure_name = MEASURE_DESCRIPTIONS.get(measure_code, measure_code)
        
        measure_maint = doc.add_paragraph()
        measure_maint.add_run(f'{measure_name}')
        measure_maint.runs[0].font.bold = True
        measure_maint.runs[0].font.size = Pt(10)
        measure_maint.runs[0].font.name = 'Calibri'
        measure_maint.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        maint_points = [
            'Follow manufacturer maintenance guidelines',
            'Conduct regular visual inspections',
            'Maintain all servicing schedules',
            'Contact RC Consultants with any concerns'
        ]
        
        for point in maint_points[:2]:  # Only 2 points to keep it concise
            point_para = doc.add_paragraph()
            point_para.add_run(f'  • {point}')
            point_para.runs[0].font.size = Pt(9)
            point_para.runs[0].font.name = 'Calibri'
            point_para.left_indent = Inches(0.25)
    
    # PAGE BREAK
    doc.add_page_break()
    
    # Add border to page 2
    add_document_border(doc.sections[-1])
    
    # PAGE 2 - CONTACT & SIGNOFF
    contact_heading = doc.add_paragraph()
    contact_heading.add_run('Contact Information')
    contact_heading.runs[0].font.size = Pt(14)
    contact_heading.runs[0].font.bold = True
    contact_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    contact_heading.runs[0].font.name = 'Calibri'
    
    contact_para = doc.add_paragraph()
    contact_para.add_run('For any questions or support regarding your retrofit measures:\n\n')
    contact_para.add_run('RC Consultants\n').bold = True
    contact_para.add_run('202 Queens Dock Business Centre\n')
    contact_para.add_run('Norfolk House, Liverpool, L1 0BG\n\n')
    contact_para.add_run('Phone: ').bold = True
    contact_para.add_run('0800 001 6127\n')
    contact_para.add_run('Email: ').bold = True
    contact_para.add_run('info@rcconsultants.co.uk\n')
    contact_para.add_run('Web: ').bold = True
    contact_para.add_run('www.rcconsultants.co.uk')
    
    for run in contact_para.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # SIGN-OFF
    signoff_heading = doc.add_paragraph()
    signoff_heading.add_run('Project Sign-Off')
    signoff_heading.runs[0].font.size = Pt(14)
    signoff_heading.runs[0].font.bold = True
    signoff_heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    signoff_heading.runs[0].font.name = 'Calibri'
    signoff_heading.space_before = Pt(20)
    
    signoff_text = doc.add_paragraph()
    signoff_text.add_run('By signing below, you confirm receipt of this handover documentation and that the installation has been completed to your satisfaction.')
    signoff_text.runs[0].font.size = Pt(10)
    signoff_text.runs[0].font.name = 'Calibri'
    signoff_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Signature table
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Light Grid Accent 1'
    
    sig_data = [
        ('Customer Signature:', ''),
        ('Retrofit Coordinator Signature:', ''),
        ('Date:', '')
    ]
    
    for i, (label, value) in enumerate(sig_data):
        row = sig_table.rows[i]
        row.height = Inches(0.7)
        
        label_cell = row.cells[0]
        label_cell.text = label
        label_para = label_cell.paragraphs[0]
        label_para.runs[0].font.bold = True
        label_para.runs[0].font.size = Pt(10)
        label_para.runs[0].font.name = 'Calibri'
        set_cell_background(label_cell, 'F2F2F2')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run('Thank you for choosing RC Consultants for your retrofit project.')
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.name = 'Calibri'
    footer.runs[0].italic = True
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc
def save_document_to_bytes(doc):
    """Save document to bytes for download"""
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io